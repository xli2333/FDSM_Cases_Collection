import os
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF

# Paths
md_path = 'website_content_structure.md'
font_path = 'font/LXGWWenKai-Regular.ttf'
pdf_path = 'website_content_structure.pdf'
docx_path = 'website_content_structure.docx'

# Read MD
try:
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
except FileNotFoundError:
    print(f"Error: {md_path} not found.")
    exit(1)

# Setup Docx
doc = Document()
# Set default font for Normal style to SimHei
style = doc.styles['Normal']
font = style.font
font.name = 'SimHei'
font.size = Pt(11) # Default size

# Define font sizes for headings
# These are just example sizes, Word will apply its default heading styles which usually include bold
# We'll override the bold/italic in the processing function for specific runs
heading1_size = Pt(20)
heading2_size = Pt(16)
heading3_size = Pt(14)

# Setup PDF
class PDF(FPDF):
    def header(self):
        pass 
    def footer(self):
        self.set_y(-15)
        try:
            self.set_font('CustomFont', '', 8)
        except:
            self.set_font('helvetica', '', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

pdf = PDF()
pdf.add_page()

# Add custom font for PDF (Crucial for Chinese)
try:
    if os.path.exists(font_path):
        pdf.add_font('CustomFont', '', font_path)
        pdf.set_font('CustomFont', '', 12)
    else:
        print(f"Warning: Font file not found at {font_path}. Chinese characters may not render in PDF.")
        pdf.set_font('helvetica', '', 12)
except Exception as e:
    print(f"Error loading font: {e}")
    pdf.set_font('helvetica', '', 12)

# Processing function
def process_line(line):
    line = line.strip()
    if not line:
        # Add a small spacer in PDF for empty lines?
        # pdf.ln(2) 
        return
    
    # --- DOCX Generation ---
    if line.startswith('# '):
        p = doc.add_paragraph()
        runner = p.add_run(line[2:])
        runner.font.name = 'SimHei'
        runner.font.size = heading1_size
        runner.bold = False
        runner.italic = False
    elif line.startswith('## '):
        p = doc.add_paragraph()
        runner = p.add_run(line[3:])
        runner.font.name = 'SimHei'
        runner.font.size = heading2_size
        runner.bold = False
        runner.italic = False
    elif line.startswith('### '):
        p = doc.add_paragraph()
        runner = p.add_run(line[4:])
        runner.font.name = 'SimHei'
        runner.font.size = heading3_size
        runner.bold = False
        runner.italic = False
    elif line.startswith('* ') or line.startswith('- '):
        p = doc.add_paragraph()
        runner = p.add_run(line[2:]) # Remove the bullet character
        runner.font.name = 'SimHei'
        runner.bold = False
        runner.italic = False
        p.paragraph_format.left_indent = Pt(20) # Add some indentation for lists
    elif line.startswith('> '):
        p = doc.add_paragraph()
        runner = p.add_run(line[2:])
        runner.font.name = 'SimHei'
        runner.bold = False
        runner.italic = False
        p.paragraph_format.left_indent = Pt(20) # Add some indentation for quotes
    else:
        p = doc.add_paragraph()
        runner = p.add_run(line)
        runner.font.name = 'SimHei'
        runner.bold = False
        runner.italic = False

    # --- PDF Generation ---
    # Reset color
    pdf.set_text_color(0, 0, 0)
    
    try:
        if line.startswith('# '):
            pdf.set_font('CustomFont', '', 20)
            pdf.ln(5)
            pdf.multi_cell(pdf.epw, 10, line[2:])
            pdf.ln(2)
        elif line.startswith('## '):
            pdf.set_font('CustomFont', '', 16)
            pdf.ln(4)
            pdf.multi_cell(pdf.epw, 8, line[3:])
            pdf.ln(2)
        elif line.startswith('### '):
            pdf.set_font('CustomFont', '', 14)
            pdf.ln(2)
            pdf.multi_cell(pdf.epw, 8, line[4:])
            pdf.ln(1)
        elif line.startswith('* ') or line.startswith('- '):
            pdf.set_font('CustomFont', '', 12)
            # Simpler bullet rendering
            pdf.multi_cell(pdf.epw, 6, '  - ' + line[2:])
        elif line.startswith('> '):
            pdf.set_font('CustomFont', '', 11)
            pdf.set_text_color(80, 80, 80) # Grey
            # Quote rendering
            pdf.multi_cell(pdf.epw, 6, '    ' + line[2:])
            pdf.set_text_color(0, 0, 0) # Reset
        else:
            pdf.set_font('CustomFont', '', 12)
            pdf.multi_cell(pdf.epw, 6, line)
            pdf.ln(1)
    except Exception as e:
        print(f"Error processing line for PDF: '{line}' - {e}")

# Run conversion
for line in lines:
    process_line(line)

# Save files
try:
    doc.save(docx_path)
    print(f"Successfully created {docx_path}")
except Exception as e:
    print(f"Error saving DOCX: {e}")

try:
    pdf.output(pdf_path)
    print(f"Successfully created {pdf_path}")
except Exception as e:
    print(f"Error saving PDF: {e}")
